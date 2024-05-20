import os
import pandas as pd
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def generate_doc(ref, name, old_salary, new_salary, old_position, new_position):
    doc = Document()

    # Adding left-aligned paragraphs for Date, REF, and Subject
    doc.add_paragraph(f"Date: January 1, 2024")
    doc.add_paragraph(f"REF:{ref}")
    doc.add_paragraph("Subject: Review and Appraisal")

    # Personalized greeting with left alignment
    doc.add_paragraph(f"Dear {name},")

    # Offer letter content
    if pd.notnull(new_position) and new_position.strip():
        content = ("We are pleased to formally notify you of changes to your employment "
                   "contract with Techkraft Inc Pvt. Ltd., effective from January 1st, 2024. "
                   "As per the recent review and assessment of your contribution, we are pleased "
                   "to inform you that the following adjustments have been made to your salary and "
                   "designation to reflect your valuable contributions and dedication to our organization.")
    else:
        content = ("We are pleased to formally notify you of changes to your employment "
                   "contract with Techkraft Inc Pvt. Ltd., effective from January 1st, 2024. "
                   "As per the recent review and assessment of your contribution, we are pleased "
                   "to inform you that the following adjustment has been made to your salary to "
                   "reflect your valuable contributions and dedication to our organization.")
    
    # Adding the content with justify alignment
    paragraph = doc.add_paragraph(content)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Add tables or other contents here, ensuring appropriate use of styles and alignment

    # Customize additional paragraphs or document elements as needed...

    # Setting the top margin of the document
    for section in doc.sections:
        section.top_margin = Cm(5)
    
    # Final paragraphs, adjusted as needed based on content and preference for alignment
    conclusion = ("It is important to note that apart from these modifications, all other terms and conditions of your "
                  "original employment contract will remain unchanged and in full effect. This encompasses your benefits, "
                  "working hours, and leave entitlements, among other provisions.")
    doc.add_paragraph(conclusion).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    doc.add_paragraph("Should you require any clarification or have any queries regarding these adjustments, please do not hesitate to reach out to the People and Culture Department.").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    doc.add_paragraph("Thank you for your ongoing dedication and contributions to Techkraft Inc Pvt. Ltd. We eagerly anticipate your continued success within the team.").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    doc.add_paragraph("Yours sincerely, \n\n\n _______________ \n Santosh Koirala \n Executive Director \n Signed Date:")

    return doc

def main():
    try:
        # Assuming you have a DataFrame or can load one with pandas as shown in your original code
        # Here, I'll simulate it with a simple example
        employee_data = pd.DataFrame({
            'ref': ['001', '002'],
            'Name': ['Alice Johnson', 'Bob Smith'],
            'Old Salary': [50000, 55000],
            'New Salary': [60000, 65000],
            'Old Position': ['Developer', 'Analyst'],
            'New Position': ['Senior Developer', 'Senior Analyst']
        })

        # Creating a folder for documents if it doesn't already exist
        folder_name = "docs"
        if not os.path.exists(folder_name):
            os.makedirs(folder_name)
        
        for index, row in employee_data.iterrows():
            # Generating the document
            doc = generate_doc(row['ref'], row['Name'], row['Old Salary'], row['New Salary'], row['Old Position'], row['New Position'])

            # Saving the document
            doc_file_name = os.path.join(folder_name, f"{row['Name'].replace(' ', '_')}_Offer_Letter.docx")
            doc.save(doc_file_name)

        print("\nOffer letters generated successfully and saved in the 'docs' folder.")
    except Exception as e:
        print("Error occurred:", e)

if __name__ == "__main__":
    main()
