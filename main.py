import os
import pandas as pd
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tkinter import Tk, Label, Button, filedialog, messagebox, StringVar, Entry, Frame, OptionMenu, Toplevel, Text
import json
from fpdf import FPDF

# Function to generate Word document for offer letter
def generate_offer_letter(ref, name, old_salary, new_salary, old_position, new_position):
    doc = Document()
    doc.styles['Normal'].paragraph_format.line_spacing = 1.3
    doc.add_paragraph(f"Date: April 10th, 2024").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f" REF:{ref} ").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph = doc.add_paragraph(" \n Subject: Review and Appraisal")
    run = paragraph.runs[0]
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"Dear {name},").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    offer_letter_paragraph = doc.add_paragraph()

    def print_line_if_new_position(new_position):
        if pd.notnull(new_position) and new_position.strip():
            run = offer_letter_paragraph.add_run("""We are pleased to formally notify you of changes to your employment contract with Techkraft Inc Pvt. Ltd., effective from January 1st, 2024. As per the recent review and assessment of your contribution, we are pleased to inform you that the following adjustments have been made to your salary and designation to reflect your valuable contributions and dedication to our organization. """)
            run.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            table = doc.add_table(rows=3, cols=3)
            table.style = 'Table Grid'
            for i, header_text in enumerate(["Particular", "Current Details", "New Details"]):
                cell = table.rows[0].cells[i]
                run = cell.paragraphs[0].add_run(header_text)
                run.bold = True
            table.rows[1].cells[0].text = "Basic Salary per month "
            table.rows[1].cells[1].text = "NRs. " + str(old_salary)
            table.rows[1].cells[2].text = "NRs. " + str(new_salary)
            table.rows[2].cells[0].text = "Designation"
            table.rows[2].cells[1].text = old_position
            table.rows[2].cells[2].text = new_position
        else:
            run = offer_letter_paragraph.add_run("""We are pleased to formally notify you of changes to your employment contract with Techkraft Inc Pvt. Ltd., effective from January 1st, 2024. As per the recent review and assessment of your contribution, we are pleased to inform you that the following adjustment have been made to your salary to reflect your valuable contributions and dedication to our organization. \n""")
            run.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            table = doc.add_table(rows=2, cols=3)
            table.style = 'Table Grid'
            for i, header_text in enumerate(["Particular", "Current Details", "New Details"]):
                cell = table.rows[0].cells[i]
                run = cell.paragraphs[0].add_run(header_text)
                run.bold = True
            table.rows[1].cells[0].text = "Salary"
            table.rows[1].cells[1].text = "NRs. " + str(old_salary)
            table.rows[1].cells[2].text = "NRs. " + str(new_salary)

    print_line_if_new_position(new_position)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(5)
    conclusion = ("\n It is important to note that apart from these modifications, all other terms and conditions of your "
                  "original employment contract will remain unchanged and in full effect. This encompasses your benefits, "
                  "working hours, and leave entitlements, among other provisions.")
    doc.add_paragraph(conclusion).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_paragraph("""\nShould you require any clarification or have any queries regarding these adjustments, please do not hesitate to reach out to the People and Culture Department. """)
    doc.add_paragraph("""\nThank you for your ongoing dedication and contributions to Techkraft Inc Pvt. Ltd. We eagerly anticipate your continued success within the team. \n""")
    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_paragraph("Yours sincerely, \n\n\n _______________ \n Santosh Koirala \n Executive Director \n Signed Date: ")
    return doc

# Function to generate PDF document for offer letter
def generate_offer_letter_pdf(ref, name, old_salary, new_salary, old_position, new_position):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Date: April 10th, 2024", ln=True, align='L')
    pdf.cell(200, 10, txt=f" REF:{ref} ", ln=True, align='L')
    pdf.set_font("Arial", 'B', size=12)
    pdf.cell(200, 10, txt="Subject: Review and Appraisal", ln=True, align='L')
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt=f"Dear {name},", ln=True, align='L')

    def print_line_if_new_position(new_position):
        if pd.notnull(new_position) and new_position.strip():
            pdf.multi_cell(0, 10, """We are pleased to formally notify you of changes to your employment contract with Techkraft Inc Pvt. Ltd., effective from January 1st, 2024. As per the recent review and assessment of your contribution, we are pleased to inform you that the following adjustments have been made to your salary and designation to reflect your valuable contributions and dedication to our organization.""")
            pdf.ln()
            pdf.set_font("Arial", 'B', size=12)
            pdf.cell(63, 10, "Particular", 1)
            pdf.cell(63, 10, "Current Details", 1)
            pdf.cell(63, 10, "New Details", 1)
            pdf.ln()
            pdf.set_font("Arial", size=12)
            pdf.cell(63, 10, "Basic Salary per month", 1)
            pdf.cell(63, 10, f"NRs. {old_salary}", 1)
            pdf.cell(63, 10, f"NRs. {new_salary}", 1)
            pdf.ln()
            pdf.cell(63, 10, "Designation", 1)
            pdf.cell(63, 10, old_position, 1)
            pdf.cell(63, 10, new_position, 1)
        else:
            pdf.multi_cell(0, 10, """We are pleased to formally notify you of changes to your employment contract with Techkraft Inc Pvt. Ltd., effective from January 1st, 2024. As per the recent review and assessment of your contribution, we are pleased to inform you that the following adjustments have been made to your salary to reflect your valuable contributions and dedication to our organization.\n""")
            pdf.ln()
            pdf.set_font("Arial", 'B', size=12)
            pdf.cell(63, 10, "Particular", 1)
            pdf.cell(63, 10, "Current Details", 1)
            pdf.cell(63, 10, "New Details", 1)
            pdf.ln()
            pdf.set_font("Arial", size=12)
            pdf.cell(63, 10, "Salary", 1)
            pdf.cell(63, 10, f"NRs. {old_salary}", 1)
            pdf.cell(63, 10, f"NRs. {new_salary}", 1)
            pdf.ln()

    print_line_if_new_position(new_position)
    conclusion = ("\nIt is important to note that apart from these modifications, all other terms and conditions of your "
                  "original employment contract will remain unchanged and in full effect. This encompasses your benefits, "
                  "working hours, and leave entitlements, among other provisions.")
    pdf.multi_cell(0, 10, conclusion)
    pdf.ln()
    pdf.multi_cell(0, 10, """\nShould you require any clarification or have any queries regarding these adjustments, please do not hesitate to reach out to the People and Culture Department. """)
    pdf.ln()
    pdf.multi_cell(0, 10, """\nThank you for your ongoing dedication and contributions to Techkraft Inc Pvt. Ltd. We eagerly anticipate your continued success within the team. \n""")
    pdf.ln(20)
    pdf.cell(200, 10, txt="Yours sincerely,", ln=True, align='L')
    pdf.cell(200, 10, txt="_______________", ln=True, align='L')
    pdf.cell(200, 10, txt="Santosh Koirala", ln=True, align='L')
    pdf.cell(200, 10, txt="Executive Director", ln=True, align='L')
    pdf.cell(200, 10, txt="Signed Date:", ln=True, align='L')
    return pdf

# Function to generate Word document for rejection letter
def generate_rejection_letter(ref, name):
    doc = Document()
    doc.styles['Normal'].paragraph_format.line_spacing = 1.3
    doc.add_paragraph(f"Date: April 10th, 2024").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f" REF:{ref} ").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph = doc.add_paragraph(" \n Subject: Application Status")
    run = paragraph.runs[0]
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"Dear {name},").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph("""We appreciate the time and effort you have invested in applying for the position at Techkraft Inc Pvt. Ltd. After careful consideration of your application and the qualifications presented, we regret to inform you that we have decided to move forward with other candidates whose experiences and qualifications more closely match our current needs.""").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph("""\nPlease do not consider this decision a reflection of your abilities. The selection process was highly competitive, and we encourage you to apply for future openings that match your skills and interests.""").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph("""\nWe thank you for your interest in joining Techkraft Inc Pvt. Ltd. and wish you all the best in your future endeavors.""").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph("Yours sincerely, \n\n\n _______________ \n HR Team \n Techkraft Inc Pvt. Ltd.")
    return doc

# Function to generate Word document for experience letter
def generate_experience_letter(ref, name, position, start_date, end_date):
    doc = Document()
    doc.styles['Normal'].paragraph_format.line_spacing = 1.3
    doc.add_paragraph(f"Date: April 10th, 2024").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f" REF:{ref} ").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph = doc.add_paragraph(" \n Subject: Experience Certificate")
    run = paragraph.runs[0]
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"To Whom It May Concern,").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph(
        f"""This is to certify that {name} was employed with Techkraft Inc Pvt. Ltd. as a {position} from {start_date} to {end_date}. During this period, {name} demonstrated outstanding professional conduct and contributed significantly to our projects."""
    ).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph(
        f"""{name} was a dedicated employee and exhibited excellent work ethics and team spirit. {name}'s skills and knowledge in the field have been a valuable asset to our company."""
    ).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph(
        f"""We wish {name} all the best in their future career endeavors."""
    ).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph("Yours sincerely, \n\n\n _______________ \n HR Team \n Techkraft Inc Pvt. Ltd.")
    return doc

# Function to generate PDF document for experience letter
def generate_experience_letter_pdf(ref, name, position, start_date, end_date):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Date: April 10th, 2024", ln=True, align='L')
    pdf.cell(200, 10, txt=f" REF:{ref} ", ln=True, align='L')
    pdf.set_font("Arial", 'B', size=12)
    pdf.cell(200, 10, txt="Subject: Experience Certificate", ln=True, align='L')
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="To Whom It May Concern,", ln=True, align='L')
    pdf.multi_cell(0, 10, f"""This is to certify that {name} was employed with Techkraft Inc Pvt. Ltd. as a {position} from {start_date} to {end_date}. During this period, {name} demonstrated outstanding professional conduct and contributed significantly to our projects.""")
    pdf.ln()
    pdf.multi_cell(0, 10, f"""{name} was a dedicated employee and exhibited excellent work ethics and team spirit. {name}'s skills and knowledge in the field have been a valuable asset to our company.""")
    pdf.ln()
    pdf.multi_cell(0, 10, f"""We wish {name} all the best in their future career endeavors.""")
    pdf.ln()
    pdf.cell(200, 10, txt="Yours sincerely,", ln=True, align='L')
    pdf.cell(200, 10, txt="_______________", ln=True, align='L')
    pdf.cell(200, 10, txt="HR Team", ln=True, align='L')
    pdf.cell(200, 10, txt="Techkraft Inc Pvt. Ltd.", ln=True, align='L')
    return pdf


# Function to generate PDF document for rejection letter
def generate_rejection_letter_pdf(ref, name):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Date: April 10th, 2024", ln=True, align='L')
    pdf.cell(200, 10, txt=f" REF:{ref} ", ln=True, align='L')
    pdf.set_font("Arial", 'B', size=12)
    pdf.cell(200, 10, txt="Subject: Application Status", ln=True, align='L')
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt=f"Dear {name},", ln=True, align='L')
    pdf.multi_cell(0, 10, """We appreciate the time and effort you have invested in applying for the position at Techkraft Inc Pvt. Ltd. After careful consideration of your application and the qualifications presented, we regret to inform you that we have decided to move forward with other candidates whose experiences and qualifications more closely match our current needs.""")
    pdf.ln()
    pdf.multi_cell(0, 10, """\nPlease do not consider this decision a reflection of your abilities. The selection process was highly competitive, and we encourage you to apply for future openings that match your skills and interests.""")
    pdf.ln()
    pdf.multi_cell(0, 10, """\nWe thank you for your interest in joining Techkraft Inc Pvt. Ltd. and wish you all the best in your future endeavors.""")
    pdf.ln()
    pdf.cell(200, 10, txt="Yours sincerely,", ln=True, align='L')
    pdf.cell(200, 10, txt="_______________", ln=True, align='L')
    pdf.cell(200, 10, txt="HR Team", ln=True, align='L')
    pdf.cell(200, 10, txt="Techkraft Inc Pvt. Ltd.", ln=True, align='L')
    return pdf


# Function to load Excel file
def load_excel_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
    excel_file_path.set(file_path)

# Function to save configurations
def save_config():
    config = {
        "excel_file": excel_file_path.get(),
        "template_type": template_type.get(),
        "output_option": output_option.get()
    }
    with open("config.json", "w") as config_file:
        json.dump(config, config_file)
    messagebox.showinfo("Success", "Configuration saved successfully.")

# Function to load configurations
def load_config():
    try:
        with open("config.json", "r") as config_file:
            config = json.load(config_file)
        excel_file_path.set(config.get("excel_file", ""))
        template_type.set(config.get("template_type", "Offer Letter"))
        output_option.set(config.get("output_option", "docx"))
        messagebox.showinfo("Success", "Configuration loaded successfully.")
    except Exception as e:
        messagebox.showerror("Error", f"Error loading configuration: {e}")

# Function to generate and preview letters
def generate_letters():
    try:
        employee_data = pd.read_excel(excel_file_path.get())
        employee_data.columns = employee_data.columns.str.strip()
        if not os.path.exists("docs"):
            os.makedirs("docs")
        for index, row in employee_data.iterrows():
            ref = row['ref']
            name = row['Name']
            if template_type.get() == "Offer Letter":
                old_salary = row['Old Salary']
                new_salary = row['New Salary']
                old_position = row['Old Position']
                new_position = row['New Position']
                if output_option.get() == "docx":
                    doc = generate_offer_letter(ref, name, old_salary, new_salary, old_position, new_position)
                    doc.save(f"docs/{name.replace(' ', '_')}_Offer_Letter.docx")
                elif output_option.get() == "pdf":
                    pdf = generate_offer_letter_pdf(ref, name, old_salary, new_salary, old_position, new_position)
                    pdf.output(f"docs/{name.replace(' ', '_')}_Offer_Letter.pdf")
            elif template_type.get() == "Rejection Letter":
                if output_option.get() == "docx":
                    doc = generate_rejection_letter(ref, name)
                    doc.save(f"docs/{name.replace(' ', '_')}_Rejection_Letter.docx")
                elif output_option.get() == "pdf":
                    pdf = generate_rejection_letter_pdf(ref, name)
                    pdf.output(f"docs/{name.replace(' ', '_')}_Rejection_Letter.pdf")
            elif template_type.get() == "Experience Letter":
                position = row['Position']
                start_date = row['Start Date']
                end_date = row['End Date']
                if output_option.get() == "docx":
                    doc = generate_experience_letter(ref, name, position, start_date, end_date)
                    doc.save(f"docs/{name.replace(' ', '_')}_Experience_Letter.docx")
                elif output_option.get() == "pdf":
                    pdf = generate_experience_letter_pdf(ref, name, position, start_date, end_date)
                    pdf.output(f"docs/{name.replace(' ', '_')}_Experience_Letter.pdf")
        messagebox.showinfo("Success", "Letters generated successfully and saved in the 'docs' folder.")
    except Exception as e:
        messagebox.showerror("Error", f"Error occurred: {e}")

# Function to preview letter
def preview_letter():
    try:
        employee_data = pd.read_excel(excel_file_path.get())
        employee_data.columns = employee_data.columns.str.strip()
        row = employee_data.iloc[0]
        ref = row['ref']
        name = row['Name']
        if template_type.get() == "Offer Letter":
            old_salary = row['Old Salary']
            new_salary = row['New Salary']
            old_position = row['Old Position']
            new_position = row['New Position']
            doc = generate_offer_letter(ref, name, old_salary, new_salary, old_position, new_position)
            preview_doc(doc)
        elif template_type.get() == "Rejection Letter":
            doc = generate_rejection_letter(ref, name)
            preview_doc(doc)
        elif template_type.get() == "Experience Letter":
            position = row['Position']
            start_date = row['Start Date']
            end_date = row['End Date']
            doc = generate_experience_letter(ref, name, position, start_date, end_date)
            preview_doc(doc)
    except Exception as e:
        messagebox.showerror("Error", f"Error occurred: {e}")

# Function to preview doc
def preview_doc(doc):
    preview_window = Toplevel()
    preview_window.title("Preview Letter")
    preview_window.geometry("800x600")
    text_widget = Text(preview_window)
    text_widget.pack(expand=True, fill="both")
    for paragraph in doc.paragraphs:
        text_widget.insert("end", paragraph.text + "\n")

# Main function to create GUI
def main():
    global excel_file_path, template_type, output_option
    root = Tk()
    root.title("Letter Generator")
    root.geometry("500x400")

    excel_file_path = StringVar()
    template_type = StringVar(value="Offer Letter")
    output_option = StringVar(value="docx")

    Label(root, text="Welcome to the Letter Generator", font=("Helvetica", 14)).pack(pady=20)
    
    frame = Frame(root)
    frame.pack(pady=10)

    Label(frame, text="Excel File:").grid(row=0, column=0, padx=5, pady=5)
    Entry(frame, textvariable=excel_file_path, width=40).grid(row=0, column=1, padx=5, pady=5)
    Button(frame, text="Browse", command=load_excel_file).grid(row=0, column=2, padx=5, pady=5)

    Label(frame, text="Template Type:").grid(row=1, column=0, padx=5, pady=5)
    OptionMenu(frame, template_type, "Offer Letter", "Rejection Letter", "Experience Letter").grid(row=1, column=1, padx=5, pady=5)

    Label(frame, text="Output Option:").grid(row=2, column=0, padx=5, pady=5)
    OptionMenu(frame, output_option, "docx", "pdf").grid(row=2, column=1, padx=5, pady=5)

    Button(root, text="Generate Letters", command=generate_letters, font=("Helvetica", 12)).pack(pady=10)
    Button(root, text="Preview Letter", command=preview_letter, font=("Helvetica", 12)).pack(pady=10)
    
    frame = Frame(root)
    frame.pack(pady=10)
    Button(frame, text="Save Configuration", command=save_config, font=("Helvetica", 12)).grid(row=0, column=0, padx=5, pady=5)
    Button(frame, text="Load Configuration", command=load_config, font=("Helvetica", 12)).grid(row=0, column=1, padx=5, pady=5)
    
    root.mainloop()

if __name__ == "__main__":
    main()
