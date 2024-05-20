import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def justify_paragraph(paragraph):
    # Split paragraph into lines
    lines = paragraph.text.split("\n")
    
    # Adjust spacing for each line except the last one
    for i in range(len(lines) - 1):
        # Calculate the number of spaces to add to each word
        words = lines[i].split()
        total_space_width = paragraph.width - sum(run.font.size for run in paragraph.runs if run.text.strip())
        total_spaces = len(words) - 1
        avg_space_width = total_space_width / total_spaces if total_spaces > 0 else 0
        
        # Adjust spacing for words in the line
        new_line = ""
        for word in words[:-1]:
            new_line += word + " " * int(avg_space_width)
        new_line += words[-1]
        
        # Replace the line with adjusted spacing
        lines[i] = new_line
    
    # Join the lines back into a paragraph
    paragraph.text = "\n".join(lines)

def generate_doc(ref, name, old_salary, new_salary, old_position, new_position):
    # Letter template
    doc = Document()
    
    # Add paragraphs with left alignment
    doc.add_paragraph(f"Date: January 1, 2024").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"\n REF:{ref} \n\n Subject: Review and Appraisal \n").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"Dear {name},\n").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    offer_letter_paragraph = doc.add_paragraph()

    def print_line_if_new_position(new_position):   
        if pd.notnull(new_position):  # Check if new_position is not NaN
            if new_position.strip():  # Check if new_position is not an empty string after stripping whitespace
                run = offer_letter_paragraph.add_run("""We are pleased to formally notify you of changes to your employment contract with Techkraft Inc Pvt. Ltd., effective from January 1st, 2024. As per the recent review and assessment of your contribution, we are pleased to inform you that the following adjustments have been made to your salary and designation to reflect your valuable contributions and dedication to our organization. \n\n""")
                run.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Justify the text

                # Add a table with 3 columns and 2 rows
                table = doc.add_table(rows=3, cols=3)  # Incremented rows for designation
                table.style = 'Table Grid'
                
                # Add column headers
                table.rows[0].cells[0].text = "Particular"
                table.rows[0].cells[1].text = "Current Details"
                table.rows[0].cells[2].text = "New Details"
                
                # Add row headers and data for Salary
                table.rows[1].cells[0].text = "Salary"
                table.rows[1].cells[1].text = str(old_salary)  # Convert old_salary to string
                table.rows[1].cells[2].text = str(new_salary)  # Convert new_salary to string
                
                # Add row headers and data for Designation
                table.rows[2].cells[0].text = "Designation"
                table.rows[2].cells[1].text = old_position
                table.rows[2].cells[2].text = new_position
            
            else:
                run = offer_letter_paragraph.add_run("""We are pleased to formally notify you of changes to your employment contract with Techkraft Inc Pvt. Ltd., effective from January 1st, 2024. As per the recent review and assessment of your contribution, we are pleased to inform you that the following adjustment have been made to your salary to reflect your valuable contributions and dedication to our organization. \n\n""")
                run.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Justify the text

                # Salary and effective date table
                table = doc.add_table(rows=2, cols=3)
                table.style = 'Table Grid'
                
                # Add column headers
                table.rows[0].cells[0].text = "Particular"
                table.rows[0].cells[1].text = "Current Details"
                table.rows[0].cells[2].text = "New Details"
                
                # Add row headers and data
                table.rows[1].cells[0].text = "Salary"
                table.rows[1].cells[1].text = str(old_salary)  # Convert old_salary to string
                table.rows[1].cells[2].text = str(new_salary)  # Convert new_salary to string
        
        else:
            run = offer_letter_paragraph.add_run("""We are pleased to formally notify you of changes to your employment contract with Techkraft Inc Pvt. Ltd., effective from January 1st, 2024. As per the recent review and assessment of your contribution, we are pleased to inform you that the following adjustment have been made to your salary to reflect your valuable contributions and dedication to our organization. \n\n""")
            run.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Justify the text

            # Salary and effective date table
            table = doc.add_table(rows=2, cols=3)
            table.style = 'Table Grid'
            
            # Add column headers
            table.rows[0].cells[0].text = "Particular"
            table.rows[0].cells[1].text = "Current Details"
            table.rows[0].cells[2].text = "New Details"
            
            # Add row headers and data
            table.rows[1].cells[0].text = "Salary"
            table.rows[1].cells[1].text = str(old_salary)  # Convert old_salary to string
            table.rows[1].cells[2].text = str(new_salary)  # Convert new_salary to string

    # Call the function to add the appropriate paragraph
    print_line_if_new_position(new_position)
    
    # Set top margin for the entire document
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(5)
        
    # Justify all paragraphs in the document
    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        justify_paragraph(paragraph)  # Adjust spacing within justified paragraphs
   
    doc.add_paragraph(""" \n It is important to note that apart from these modifications, all other terms and conditions of your original employment contract will remain unchanged and in full effect. This encompasses your benefits, working hours, and leave entitlements, among other provisions.""")
    doc.add_paragraph("""\n Should you require any clarification or have any queries regarding these adjustments, please do not hesitate to reach out to the People and Culture Department. """)
    
    doc.add_paragraph("""\n Thank you for your ongoing dedication and contributions to Techkraft Inc Pvt. Ltd. We eagerly anticipate your continued success within the team. \n""")
    
    doc.add_paragraph("Yours sincerely, \n\n\n _______________ \n Santosh Koirala \n Executive Director \n Signed Date:	")

    return doc

def main():
    try:
        # Load employee data from Excel sheet
        employee_data = pd.read_excel('Booktest.xlsx')
        print("Successfully loaded data from the Excel sheet:")
        print(employee_data.head())  # Print the first few rows of the DataFrame
        
        # Remove leading or trailing whitespaces from column names
        employee_data.columns = employee_data.columns.str.strip()
        
        print("Columns in the DataFrame:", employee_data.columns)
        
        # Create a folder to save offer letters if it doesn't exist
        folder_name = "docs"
        if not os.path.exists(folder_name):
            os.makedirs(folder_name)
        
        for index, row in employee_data.iterrows():
            # Extracting employee's details
            ref = row['ref']
            name = row['Name']
            old_salary = row['Old Salary']
            new_salary = row['New Salary']
            old_position = row['Old Position']
            new_position = row['New Position']

            # Generate offer letter
            doc = generate_doc(ref, name, old_salary, new_salary, old_position, new_position)
            
            # Save offer letter to a file
            doc.save(f"{name.replace(' ', '_')}_Offer_Letter_test1.docx")
        
        print("\nOffer letters generated successfully and saved in the 'docs' folder.")
    except Exception as e:
        print("Error occurred:", e)

if __name__ == "__main__":
    main()
